# utils/word_export.py
from typing import Dict, Any
from io import BytesIO

from docx import Document


def build_meeting_docx(meeting_json: Dict[str, Any], summary: str) -> BytesIO:
    """
    根据结构化 JSON 纪要和摘要构建一个 Word 文档，并返回 BytesIO 对象。
    结构示例：
    {
        "Topic": "...",
        "KeyPoints": ["...", "..."],
        "Decisions": ["...", "..."],
        "ActionItems": [
            {"Who": "张三", "What": "完成原型设计", "When": "本周五"},
            ...
        ]
    }
    """
    doc = Document()
    doc.add_heading("会议纪要 MeetingAgent Pro", level=1)

    # 摘要
    doc.add_heading("一、会议摘要", level=2)
    doc.add_paragraph(summary or "（暂无摘要）")

    # 议题
    topic = meeting_json.get("Topic") or "（未提取）"
    doc.add_heading("二、会议议题", level = 2)
    doc.add_paragraph(topic)

    # 关键讨论点
    doc.add_heading("三、关键讨论点", level=2)
    key_points = meeting_json.get("KeyPoints") or []
    if isinstance(key_points, list):
        for kp in key_points:
            doc.add_paragraph(kp, style="List Bullet")
    else:
        doc.add_paragraph(str(key_points))

    # 决策结论
    doc.add_heading("四、决策结论", level=2)
    decisions = meeting_json.get("Decisions") or []
    if isinstance(decisions, list):
        for dec in decisions:
            doc.add_paragraph(dec, style="List Number")
    else:
        doc.add_paragraph(str(decisions))

    # 待办事项表格
    doc.add_heading("五、待办事项（Action Items）", level=2)
    action_items = meeting_json.get("ActionItems") or []

    if isinstance(action_items, list) and action_items:
        table = doc.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "责任人 (Who)"
        hdr_cells[1].text = "任务内容 (What)"
        hdr_cells[2].text = "截止时间 (When)"

        for item in action_items:
            row_cells = table.add_row().cells
            row_cells[0].text = str(item.get("Who", ""))
            row_cells[1].text = str(item.get("What", ""))
            row_cells[2].text = str(item.get("When", ""))
    else:
        doc.add_paragraph("暂无待办事项。")

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer